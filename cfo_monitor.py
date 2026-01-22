#!/usr/bin/env python3
"""
CFO Change Monitor - Enhanced Version with Tear Sheets
Monitors SEC EDGAR filings and business news
Generates detailed company and individual tear sheets as Word documents
100% free and ToS compliant
"""

import requests
import feedparser
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime, timedelta
from io import BytesIO
import time
import os
import json
import re

class CFOMonitor:
    def __init__(self, email_to, email_from, email_password, anthropic_api_key=None):
        self.email_to = email_to
        self.email_from = email_from
        self.email_password = email_password
        self.anthropic_api_key = anthropic_api_key
        self.results = []
        self.tear_sheets = []
        
    def search_sec_filings(self):
        """Search SEC EDGAR for 8-K filings mentioning CFO changes"""
        print("Searching SEC EDGAR filings...")
        
        headers = {
            'User-Agent': f'{self.email_from} CFO Monitor Script'
        }
        
        base_url = "https://www.sec.gov/cgi-bin/browse-edgar"
        params = {
            'action': 'getcurrent',
            'type': '8-K',
            'count': '100',
            'output': 'atom'
        }
        
        try:
            response = requests.get(base_url, params=params, headers=headers, timeout=30)
            time.sleep(0.1)
            
            if response.status_code == 200:
                feed = feedparser.parse(response.content)
                
                for entry in feed.entries:
                    title = entry.get('title', '')
                    summary = entry.get('summary', '')
                    link = entry.get('link', '')
                    
                    text = (title + ' ' + summary).lower()
                    cfo_keywords = ['cfo', 'chief financial officer', 'financial officer']
                    
                    if any(keyword in text for keyword in cfo_keywords):
                        company_name = title.split('(')[0].strip() if '(' in title else title
                        
                        result = {
                            'source': 'SEC EDGAR',
                            'company': company_name,
                            'title': title,
                            'summary': summary[:300],
                            'url': link,
                            'date': entry.get('published', 'N/A'),
                            'individual': self._extract_individual_name(title, summary)
                        }
                        
                        self.results.append(result)
                        
            print(f"Found {len([r for r in self.results if r['source'] == 'SEC EDGAR'])} SEC filings")
            
        except Exception as e:
            print(f"Error searching SEC: {e}")
    
    def search_news(self):
        """Search Google News RSS feeds for CFO appointment news"""
        print("Searching business news...")
        
        search_queries = [
            'CFO appointed',
            'CFO hired',
            'Chief Financial Officer joins',
            'CFO departure',
            'CFO steps down',
            'new CFO',
            'names CFO',
            'appoints Chief Financial Officer'
        ]
        
        for query in search_queries:
            try:
                url = f"https://news.google.com/rss/search?q={query.replace(' ', '+')}&hl=en-US&gl=US&ceid=US:en"
                feed = feedparser.parse(url)
                time.sleep(0.5)
                
                for entry in feed.entries[:5]:
                    pub_date = entry.get('published_parsed')
                    if pub_date:
                        entry_date = datetime(*pub_date[:6])
                        if datetime.now() - entry_date > timedelta(days=2):
                            continue
                    
                    title = entry.get('title', '')
                    link = entry.get('link', '')
                    source = entry.get('source', {}).get('title', 'Unknown')
                    
                    if not any(r.get('url') == link for r in self.results):
                        result = {
                            'source': f'News: {source}',
                            'company': self._extract_company(title),
                            'title': title,
                            'summary': '',
                            'url': link,
                            'date': entry.get('published', 'N/A'),
                            'individual': self._extract_individual_name(title, '')
                        }
                        
                        self.results.append(result)
                        
            except Exception as e:
                print(f"Error searching news for '{query}': {e}")
        
        news_count = len([r for r in self.results if r['source'].startswith('News')])
        print(f"Found {news_count} news articles")
    
    def _extract_company(self, title):
        """Extract company name from news title"""
        keywords = ['appoints', 'hires', 'names', 'announces', 'welcomes']
        for keyword in keywords:
            if keyword in title.lower():
                parts = title.lower().split(keyword)
                if parts[0]:
                    return parts[0].strip().title()
        return 'Company in article'
    
    def _extract_individual_name(self, title, summary):
        """Extract individual's name from title or summary"""
        text = title + ' ' + summary
        
        # Common patterns for names in CFO announcements
        patterns = [
            r'(?:appoints|names|hires|welcomes)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\s+(?:as|named|appointed|joins)',
            r'CFO\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                name = match.group(1).strip()
                # Filter out common false positives
                if len(name.split()) >= 2 and name not in ['Chief Financial', 'Financial Officer']:
                    return name
        
        return None
    
    def generate_tear_sheets(self):
        """Generate tear sheets for companies and individuals using Claude API"""
        if not self.anthropic_api_key:
            print("Skipping tear sheet generation - no API key provided")
            return
        
        print(f"\nGenerating tear sheets for {len(self.results)} findings...")
        
        for idx, result in enumerate(self.results):
            print(f"Processing {idx + 1}/{len(self.results)}: {result['company']}")
            
            # Generate company tear sheet
            company_doc = self._generate_company_tear_sheet(result)
            if company_doc:
                self.tear_sheets.append({
                    'type': 'company',
                    'company': result['company'],
                    'document': company_doc,
                    'filename': f"{self._sanitize_filename(result['company'])}_Company_TearSheet.docx"
                })
            
            # Generate individual tear sheet if we have a name
            if result.get('individual'):
                individual_doc = self._generate_individual_tear_sheet(result)
                if individual_doc:
                    self.tear_sheets.append({
                        'type': 'individual',
                        'individual': result['individual'],
                        'company': result['company'],
                        'document': individual_doc,
                        'filename': f"{self._sanitize_filename(result['individual'])}_Individual_TearSheet.docx"
                    })
            
            time.sleep(2)  # Rate limiting for API calls
        
        print(f"Generated {len(self.tear_sheets)} tear sheet documents")
    
    def _generate_company_tear_sheet(self, result):
        """Generate company tear sheet using Claude API"""
        try:
            # Search for company information
            company_name = result['company']
            
            prompt = f"""Research and create a detailed company tear sheet for {company_name}. 

Use this structure:

**COMPANY TEAR SHEET: {company_name}**
Generated: {datetime.now().strftime("%B %d, %Y")}

**Section 1: Company Overview**
- Generate background bullets about the company with focus on industry/sector, sources of income, headquarters, founding year, employees, and relevant background
- Provide competitive landscape summary
- Extract from recent 10-K filings: revenue by business unit (%), operating/SG&A expenses (%), total revenue/costs/operating income, year-over-year changes

**Section 2: Leadership**
- Summary of executive leadership team (C-suite) with LinkedIn profiles
- Board of Directors summary

**Section 3: Company News and Strategic Initiatives (Last 5 Years)**
- Key strategic initiatives grouped by year
- Focus on M&A activity, funding, new products/services, leadership changes
- Include sources/links

**Section 4: SWOT Analysis**
- 5-10 bullets each for: Strengths, Weaknesses, Opportunities (external), Threats (external)

**Section 5: Locations**
- List headquarters, regional offices, manufacturing sites, and other locations

Provide comprehensive, well-researched information in plain text format with simple formatting."""

            response = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": self.anthropic_api_key,
                    "anthropic-version": "2023-06-01"
                },
                json={
                    "model": "claude-sonnet-4-20250514",
                    "max_tokens": 4000,
                    "messages": [{"role": "user", "content": prompt}]
                },
                timeout=60
            )
            
            if response.status_code == 200:
                data = response.json()
                content = data.get('content', [])
                if content and content[0].get('type') == 'text':
                    return content[0].get('text', '')
            
        except Exception as e:
            print(f"Error generating company tear sheet: {e}")
        
        return None
    
    def _generate_individual_tear_sheet(self, result):
        """Generate individual tear sheet using Claude API"""
        try:
            individual_name = result.get('individual')
            company_name = result['company']
            
            prompt = f"""Research and create a detailed individual tear sheet for {individual_name}, the new CFO at {company_name}.

Use this structure:

**INDIVIDUAL TEAR SHEET: {individual_name}**
Generated: {datetime.now().strftime("%B %d, %Y")}

**Section 1: Executive Overview**
- Full name
- Current title and company
- Industry/sector focus
- Primary responsibilities
- Years in role and total leadership experience
- Location
- Education (degrees, institutions, certifications)
- Board/advisory roles
- Awards/recognition
- LinkedIn and public presence summary

**Section 2: Leadership Team & Board Connectivity**
- Role within C-suite
- Key peer relationships
- Organizational structure insights
- Board exposure and involvement

**Section 3: Professional Milestones & Strategic Initiatives (Last 5-10 Years)**
Grouped by year:
- M&A activity
- Capital structure moves
- Major operational initiatives
- Strategic pivots
- Leadership changes
- Quantified achievements

**Section 4: SWOT Analysis (Individual)**
- Strengths (5 bullets): expertise, experience, capabilities
- Weaknesses (5 bullets): gaps or limitations
- Opportunities (5 bullets - external): career growth, market trends
- Threats (5 bullets - external): market challenges, competition

**Section 5: Location & Mobility**
- Primary residence (city, state)
- Current work locations

Provide comprehensive, well-researched information in plain text format."""

            response = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": self.anthropic_api_key,
                    "anthropic-version": "2023-06-01"
                },
                json={
                    "model": "claude-sonnet-4-20250514",
                    "max_tokens": 4000,
                    "messages": [{"role": "user", "content": prompt}]
                },
                timeout=60
            )
            
            if response.status_code == 200:
                data = response.json()
                content = data.get('content', [])
                if content and content[0].get('type') == 'text':
                    return content[0].get('text', '')
            
        except Exception as e:
            print(f"Error generating individual tear sheet: {e}")
        
        return None
    
    def _sanitize_filename(self, name):
        """Sanitize filename by removing invalid characters"""
        return re.sub(r'[^\w\s-]', '', name).strip().replace(' ', '_')[:50]
    
    def _create_word_document(self, content):
        """Create a simple Word document from text content"""
        # Using python-docx library to create proper Word documents
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Set narrow margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Parse and add content
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                # Headers (lines starting with **)
                if line.startswith('**') and line.endswith('**'):
                    text = line.strip('*').strip()
                    para = doc.add_paragraph(text)
                    para.style = 'Heading 1'
                    run = para.runs[0]
                    run.font.size = Pt(14)
                    run.bold = True
                
                # Bullet points
                elif line.startswith('-') or line.startswith('•'):
                    text = line.lstrip('-•').strip()
                    para = doc.add_paragraph(text, style='List Bullet')
                    run = para.runs[0]
                    run.font.size = Pt(11)
                
                # Regular paragraphs
                else:
                    para = doc.add_paragraph(line)
                    run = para.runs[0]
                    run.font.size = Pt(11)
            
            # Save to BytesIO
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)
            return docx_file
            
        except ImportError:
            print("Warning: python-docx not installed. Creating text file instead.")
            # Fallback to text file
            text_file = BytesIO(content.encode('utf-8'))
            return text_file
    
    def send_email(self):
        """Send email digest with all findings and tear sheet attachments"""
        if not self.results:
            print("No CFO changes found today - no email sent")
            return
        
        print(f"\nPreparing email with {len(self.results)} findings and {len(self.tear_sheets)} attachments...")
        
        # Create multipart message
        msg = MIMEMultipart()
        msg['Subject'] = f'CFO Changes Alert - {datetime.now().strftime("%B %d, %Y")} ({len(self.results)} findings)'
        msg['From'] = self.email_from
        msg['To'] = self.email_to
        
        # Create email body with 4-sentence summary
        body = self._create_email_body()
        msg.attach(MIMEText(body, 'html'))
        
        # Attach tear sheets as Word documents
        for tear_sheet in self.tear_sheets:
            try:
                docx_file = self._create_word_document(tear_sheet['document'])
                
                part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
                part.set_payload(docx_file.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{tear_sheet["filename"]}"')
                msg.attach(part)
                
            except Exception as e:
                print(f"Error attaching {tear_sheet['filename']}: {e}")
        
        try:
            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
                server.login(self.email_from, self.email_password)
                server.send_message(msg)
            
            print("✓ Email sent successfully with attachments!")
            
        except Exception as e:
            print(f"Error sending email: {e}")
    
    def _create_email_body(self):
        """Create concise 4-sentence email body with summary"""
        
        num_companies = len(set(r['company'] for r in self.results))
        num_individuals = len([r for r in self.results if r.get('individual')])
        num_sec = len([r for r in self.results if r['source'] == 'SEC EDGAR'])
        num_news = len([r for r in self.results if r['source'].startswith('News')])
        
        summary = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                .summary {{ background-color: #ecf0f1; padding: 20px; margin: 20px 0; border-radius: 5px; }}
                .highlight {{ color: #2c3e50; font-weight: bold; }}
                ul {{ margin: 10px 0; }}
                li {{ margin: 5px 0; }}
            </style>
        </head>
        <body>
            <div class="summary">
                <h2>CFO Changes Summary - {datetime.now().strftime("%B %d, %Y")}</h2>
                <p>Our automated monitoring system has identified <span class="highlight">{len(self.results)} CFO-related changes</span> across <span class="highlight">{num_companies} companies</span> in the past 24 hours. We found {num_sec} official SEC filings and {num_news} news articles reporting these executive movements. Detailed company and individual tear sheets are attached as Word documents for your review and analysis. Please see the attachments for comprehensive research on each identified change.</p>
                
                <h3>Quick Stats:</h3>
                <ul>
                    <li><strong>Total Findings:</strong> {len(self.results)}</li>
                    <li><strong>Companies Affected:</strong> {num_companies}</li>
                    <li><strong>SEC Filings:</strong> {num_sec}</li>
                    <li><strong>News Articles:</strong> {num_news}</li>
                    <li><strong>Tear Sheets Generated:</strong> {len(self.tear_sheets)}</li>
                </ul>
                
                <h3>Identified Changes:</h3>
                <ul>
        """
        
        for result in self.results:
            individual = result.get('individual', 'Individual name not identified')
            summary += f"<li><strong>{result['company']}</strong> - {individual}</li>"
        
        summary += """
                </ul>
            </div>
        </body>
        </html>
        """
        
        return summary
    
    def run(self):
        """Main execution method"""
        print("=" * 70)
        print("CFO Change Monitor - Enhanced Edition with Tear Sheets")
        print("=" * 70)
        
        self.search_sec_filings()
        self.search_news()
        
        if self.results:
            self.generate_tear_sheets()
        
        self.send_email()
        
        print("=" * 70)
        print("Scan complete!")
        print("=" * 70)


if __name__ == "__main__":
    EMAIL_TO = os.environ.get('EMAIL_TO', 'your-email@gmail.com')
    EMAIL_FROM = os.environ.get('EMAIL_FROM', 'your-email@gmail.com')
    EMAIL_PASSWORD = os.environ.get('EMAIL_PASSWORD', 'your-app-password')
    ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', None)
    
    monitor = CFOMonitor(
        email_to=EMAIL_TO,
        email_from=EMAIL_FROM,
        email_password=EMAIL_PASSWORD,
        anthropic_api_key=ANTHROPIC_API_KEY
    )
    
    monitor.run()
